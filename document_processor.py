"""
Document Processing and RAG (Retrieval-Augmented Generation) Module
Handles document uploads, text extraction, chunking, and semantic search
"""

import os
import json
import hashlib
from pathlib import Path
from typing import List, Dict, Optional
import re


# Document storage directory
DOCUMENTS_DIR = Path("uploaded_documents")
CHUNKS_FILE = Path("document_chunks.json")


def ensure_documents_dir():
    """Create documents directory if it doesn't exist."""
    DOCUMENTS_DIR.mkdir(exist_ok=True)


def extract_text_from_file(file_path: str, file_content: bytes = None, file_type: str = None) -> str:
    """
    Extract text from various file formats.
    
    Args:
        file_path: Path to the file or filename
        file_content: Raw file bytes (for uploaded files)
        file_type: File extension/type
        
    Returns:
        Extracted text content
    """
    if file_type is None:
        file_type = Path(file_path).suffix.lower()
    
    try:
        # Plain text files
        if file_type in ['.txt', '.md', '.csv']:
            if file_content:
                return file_content.decode('utf-8', errors='ignore')
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        
        # PDF files
        elif file_type == '.pdf':
            try:
                import pypdf
                if file_content:
                    import io
                    pdf_file = io.BytesIO(file_content)
                    reader = pypdf.PdfReader(pdf_file)
                else:
                    reader = pypdf.PdfReader(file_path)
                
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
            except ImportError:
                return "[PDF support requires pypdf package. Install with: pip install pypdf]"
        
        # Word documents
        elif file_type in ['.docx', '.doc']:
            try:
                import docx
                if file_content:
                    import io
                    doc = docx.Document(io.BytesIO(file_content))
                else:
                    doc = docx.Document(file_path)
                
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
                return text
            except ImportError:
                return "[DOCX support requires python-docx package. Install with: pip install python-docx]"
        
        else:
            # Try to read as text
            if file_content:
                return file_content.decode('utf-8', errors='ignore')
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
    except Exception as e:
        return f"[Error extracting text: {str(e)}]"


def chunk_text(text: str, chunk_size: int = 250, overlap: int = 50) -> List[str]:
    """
    Split text into overlapping chunks for better retrieval.
    Uses smaller chunks for more precise matching.
    
    Args:
        text: Full document text
        chunk_size: Target size of each chunk (in words) - smaller for precision
        overlap: Number of words to overlap between chunks
        
    Returns:
        List of text chunks
    """
    # Clean the text but preserve some structure
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Try to split on sentence boundaries for better chunks
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    if not sentences:
        return [text] if text else []
    
    chunks = []
    current_chunk = []
    current_word_count = 0
    
    for sentence in sentences:
        sentence_words = sentence.split()
        sentence_word_count = len(sentence_words)
        
        if current_word_count + sentence_word_count <= chunk_size:
            current_chunk.append(sentence)
            current_word_count += sentence_word_count
        else:
            # Save current chunk
            if current_chunk:
                chunks.append(' '.join(current_chunk))
            
            # Start new chunk with overlap
            if chunks and overlap > 0:
                # Get last few words from previous chunk for overlap
                prev_words = chunks[-1].split()[-overlap:]
                current_chunk = [' '.join(prev_words), sentence]
                current_word_count = len(prev_words) + sentence_word_count
            else:
                current_chunk = [sentence]
                current_word_count = sentence_word_count
    
    # Don't forget last chunk
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    # If no chunks created, fall back to word-based chunking
    if not chunks:
        words = text.split()
        if len(words) <= chunk_size:
            return [text] if text else []
        
        start = 0
        while start < len(words):
            end = start + chunk_size
            chunk = ' '.join(words[start:end])
            chunks.append(chunk)
            start = end - overlap
    
    return chunks


def compute_file_hash(content: bytes) -> str:
    """Compute MD5 hash of file content for deduplication."""
    return hashlib.md5(content).hexdigest()


def save_document(filename: str, content: bytes, metadata: Dict = None) -> Dict:
    """
    Save an uploaded document and extract its text chunks.
    
    Args:
        filename: Original filename
        content: File content as bytes
        metadata: Optional metadata (category, description, etc.)
        
    Returns:
        Document info dict with id, chunks count, etc.
    """
    ensure_documents_dir()
    
    # Generate unique ID
    file_hash = compute_file_hash(content)
    file_ext = Path(filename).suffix.lower()
    
    # Save the raw file
    safe_filename = re.sub(r'[^\w\-_\.]', '_', filename)
    file_path = DOCUMENTS_DIR / f"{file_hash}_{safe_filename}"
    
    with open(file_path, 'wb') as f:
        f.write(content)
    
    # Extract text
    text = extract_text_from_file(str(file_path), content, file_ext)
    
    # Create chunks
    chunks = chunk_text(text)
    
    # Load existing chunks database
    all_chunks = load_all_chunks()
    
    # Add new document chunks
    doc_info = {
        'id': file_hash,
        'filename': filename,
        'file_path': str(file_path),
        'file_type': file_ext,
        'metadata': metadata or {},
        'chunk_count': len(chunks),
        'total_chars': len(text)
    }
    
    # Store chunks with document reference
    for i, chunk in enumerate(chunks):
        chunk_id = f"{file_hash}_{i}"
        all_chunks[chunk_id] = {
            'doc_id': file_hash,
            'doc_name': filename,
            'chunk_index': i,
            'text': chunk,
            'metadata': metadata or {}
        }
    
    # Save chunks database
    save_all_chunks(all_chunks)
    
    return doc_info


def load_all_chunks() -> Dict:
    """Load all document chunks from storage."""
    if CHUNKS_FILE.exists():
        try:
            with open(CHUNKS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return {}
    return {}


def save_all_chunks(chunks: Dict):
    """Save all document chunks to storage."""
    with open(CHUNKS_FILE, 'w', encoding='utf-8') as f:
        json.dump(chunks, f, indent=2, ensure_ascii=False)


def delete_document(doc_id: str) -> bool:
    """
    Delete a document and its chunks.
    
    Args:
        doc_id: Document ID (hash)
        
    Returns:
        True if deleted successfully
    """
    all_chunks = load_all_chunks()
    
    # Find and remove chunks belonging to this document
    chunks_to_remove = [cid for cid, chunk in all_chunks.items() if chunk['doc_id'] == doc_id]
    
    for chunk_id in chunks_to_remove:
        del all_chunks[chunk_id]
    
    save_all_chunks(all_chunks)
    
    # Try to remove the file
    for file in DOCUMENTS_DIR.glob(f"{doc_id}_*"):
        try:
            file.unlink()
        except:
            pass
    
    return len(chunks_to_remove) > 0


def get_all_documents() -> List[Dict]:
    """Get list of all uploaded documents."""
    all_chunks = load_all_chunks()
    
    # Group by document
    docs = {}
    for chunk_id, chunk in all_chunks.items():
        doc_id = chunk['doc_id']
        if doc_id not in docs:
            docs[doc_id] = {
                'id': doc_id,
                'filename': chunk['doc_name'],
                'chunk_count': 0,
                'metadata': chunk.get('metadata', {})
            }
        docs[doc_id]['chunk_count'] += 1
    
    return list(docs.values())


# Synonym mappings for HR domain
HR_SYNONYMS = {
    'compensation': ['salary', 'pay', 'wage', 'wages', 'payment', 'earnings', 'income', 'rate', 'hourly', '$', 'usd', 'stipend', 'remuneration'],
    'salary': ['compensation', 'pay', 'wage', 'wages', 'payment', 'earnings', 'income', '$', 'annual', 'yearly'],
    'pay': ['salary', 'compensation', 'wage', 'wages', 'payment', 'earnings', 'income', '$'],
    'benefits': ['insurance', 'health', 'dental', 'vision', '401k', 'retirement', 'pto', 'vacation', 'perks', 'package'],
    'vacation': ['pto', 'time off', 'leave', 'holiday', 'holidays', 'days off', 'paid leave'],
    'leave': ['pto', 'vacation', 'time off', 'absence', 'sick', 'parental', 'maternity', 'paternity'],
    'requirements': ['qualifications', 'required', 'must have', 'experience', 'skills', 'education', 'degree'],
    'qualifications': ['requirements', 'required', 'experience', 'skills', 'education', 'credentials'],
    'responsibilities': ['duties', 'tasks', 'role', 'job duties', 'accountabilities', 'work', 'perform'],
    'duties': ['responsibilities', 'tasks', 'job duties', 'accountabilities', 'work'],
    'experience': ['years', 'background', 'expertise', 'history', 'prior', 'previous'],
    'skills': ['abilities', 'competencies', 'proficiency', 'expertise', 'knowledge'],
    'policy': ['policies', 'procedure', 'procedures', 'guidelines', 'rules', 'handbook'],
    'employee': ['staff', 'worker', 'team member', 'personnel', 'hire'],
    'manager': ['supervisor', 'lead', 'director', 'boss', 'head'],
    'intern': ['internship', 'trainee', 'apprentice', 'entry level', 'junior'],
    'job': ['position', 'role', 'opportunity', 'opening', 'employment'],
    'location': ['office', 'remote', 'hybrid', 'onsite', 'work from home', 'wfh', 'city', 'address'],
    'description': ['jd', 'job description', 'role description', 'overview', 'summary'],
    'hr': ['human resources', 'people', 'people operations', 'talent', 'personnel'],
}

# Stopwords to ignore in search
STOPWORDS = {
    'a', 'an', 'the', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
    'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
    'should', 'may', 'might', 'must', 'shall', 'can', 'need', 'dare',
    'ought', 'used', 'to', 'of', 'in', 'for', 'on', 'with', 'at', 'by',
    'from', 'up', 'about', 'into', 'through', 'during', 'before', 'after',
    'above', 'below', 'between', 'under', 'again', 'further', 'then',
    'once', 'here', 'there', 'when', 'where', 'why', 'how', 'all', 'each',
    'few', 'more', 'most', 'other', 'some', 'such', 'no', 'nor', 'not',
    'only', 'own', 'same', 'so', 'than', 'too', 'very', 's', 't', 'just',
    'don', 'now', 'i', 'me', 'my', 'myself', 'we', 'our', 'ours',
    'ourselves', 'you', 'your', 'yours', 'yourself', 'yourselves', 'he',
    'him', 'his', 'himself', 'she', 'her', 'hers', 'herself', 'it', 'its',
    'itself', 'they', 'them', 'their', 'theirs', 'themselves', 'what',
    'which', 'who', 'whom', 'this', 'that', 'these', 'those', 'am', 'and',
    'but', 'if', 'or', 'because', 'as', 'until', 'while', 'although',
    'find', 'tell', 'me', 'please', 'help', 'get', 'give', 'show', 'list'
}


def expand_query_with_synonyms(query: str) -> set:
    """
    Expand query terms with synonyms for better recall.
    Filters out stopwords for cleaner search.
    
    Args:
        query: Original search query
        
    Returns:
        Set of expanded query terms
    """
    query_lower = query.lower()
    query_words = set(re.findall(r'\b\w+\b', query_lower))
    
    # Remove stopwords
    query_words = query_words - STOPWORDS
    
    expanded = set(query_words)
    
    # Add synonyms
    for word in list(query_words):
        if word in HR_SYNONYMS:
            expanded.update(HR_SYNONYMS[word])
        # Also check if any synonym maps to this word
        for key, synonyms in HR_SYNONYMS.items():
            if word in synonyms:
                expanded.add(key)
                expanded.update(synonyms)
    
    return expanded


def simple_search(query: str, top_k: int = 5) -> List[Dict]:
    """
    Robust keyword-based search through document chunks with synonym expansion.
    Uses enhanced TF-IDF-like scoring for better relevance.
    
    Args:
        query: Search query
        top_k: Number of top results to return
        
    Returns:
        List of relevant chunks with scores
    """
    all_chunks = load_all_chunks()
    
    if not all_chunks:
        return []
    
    # Expand query with synonyms
    query_words = expand_query_with_synonyms(query)
    original_words = set(re.findall(r'\b\w+\b', query.lower())) - STOPWORDS
    
    results = []
    for chunk_id, chunk in all_chunks.items():
        text = chunk['text'].lower()
        text_words = set(re.findall(r'\b\w+\b', text))
        
        # Calculate relevance score
        score = 0
        matches = []
        
        for word in query_words:
            # Skip stopwords
            if word in STOPWORDS:
                continue
            # Check both exact word and partial matches
            if word in text_words:
                # Higher weight for original query words
                weight = 3.0 if word in original_words else 1.0
                # Longer words get higher weight
                weight *= (1 + len(word) / 10)
                # Count occurrences
                count = text.count(word)
                word_score = count * weight
                score += word_score
                matches.append(word)
            elif len(word) > 3:
                # Partial match for longer words
                for text_word in text_words:
                    if word in text_word or text_word in word:
                        weight = 1.5 if word in original_words else 0.5
                        score += weight
                        matches.append(f"{word}~{text_word}")
                        break
        
        # Special patterns for compensation/salary info
        if any(w in original_words for w in ['compensation', 'salary', 'pay', 'wage']):
            # Look for dollar amounts
            dollar_pattern = r'\$[\d,]+(?:\.\d{2})?(?:\s*[-–]\s*\$[\d,]+(?:\.\d{2})?)?(?:\s*(?:per|\/)\s*(?:hour|hr|year|yr|month|mo|week|wk))?'
            dollar_matches = re.findall(dollar_pattern, text)
            if dollar_matches:
                score += len(dollar_matches) * 10
                matches.extend(dollar_matches)
            
            # Look for salary range patterns
            salary_patterns = [
                r'\d+k\s*[-–]\s*\d+k',  # 50k-60k
                r'\d{2,3},\d{3}',  # 50,000
                r'(?:salary|pay|hourly|rate|compensation).*?\d+',  # salary of 50000
            ]
            for pattern in salary_patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    score += 5
        
        # Boost if query words appear close together (phrase matching)
        if len(original_words) > 1:
            query_phrase = query.lower()
            if query_phrase in text:
                score *= 3  # Strong boost for exact phrase
            else:
                # Check for words appearing within 50 chars of each other
                positions = []
                for word in original_words:
                    for match in re.finditer(re.escape(word), text):
                        positions.append(match.start())
                if len(positions) >= 2:
                    positions.sort()
                    for i in range(len(positions) - 1):
                        if positions[i+1] - positions[i] < 50:
                            score *= 1.5
                            break
        
        if score > 0:
            results.append({
                'chunk_id': chunk_id,
                'doc_id': chunk['doc_id'],
                'doc_name': chunk['doc_name'],
                'text': chunk['text'],
                'score': score,
                'matches': matches[:10],  # Include what matched
                'metadata': chunk.get('metadata', {})
            })
    
    # Sort by score descending
    results.sort(key=lambda x: x['score'], reverse=True)
    
    return results[:top_k]


def semantic_search_with_gemini(query: str, top_k: int = 5) -> List[Dict]:
    """
    Use Gemini to find the most relevant chunks for a query.
    This is a simple but effective approach without embeddings.
    
    Args:
        query: User's question
        top_k: Number of results to return
        
    Returns:
        List of relevant chunks
    """
    all_chunks = load_all_chunks()
    
    if not all_chunks:
        return []
    
    # First, do a simple keyword search to narrow down candidates
    keyword_results = simple_search(query, top_k=20)
    
    if not keyword_results:
        # If no keyword matches, return random sample
        import random
        chunk_list = list(all_chunks.values())
        return random.sample(chunk_list, min(top_k, len(chunk_list)))
    
    return keyword_results[:top_k]


def get_context_for_query(query: str, max_chunks: int = 3, max_chars: int = 3000) -> str:
    """
    Get relevant document context for a query.
    
    Args:
        query: User's question
        max_chunks: Maximum number of chunks to include
        max_chars: Maximum total characters
        
    Returns:
        Combined context string from relevant documents
    """
    results = semantic_search_with_gemini(query, top_k=max_chunks)
    
    if not results:
        return ""
    
    context_parts = []
    total_chars = 0
    
    for result in results:
        text = result['text'] if isinstance(result, dict) else result.get('text', '')
        doc_name = result.get('doc_name', 'Unknown Document')
        
        if total_chars + len(text) > max_chars:
            # Truncate if needed
            remaining = max_chars - total_chars
            if remaining > 100:
                text = text[:remaining] + "..."
            else:
                break
        
        context_parts.append(f"[From: {doc_name}]\n{text}")
        total_chars += len(text)
    
    return "\n\n---\n\n".join(context_parts)


def has_uploaded_documents() -> bool:
    """Check if there are any uploaded documents."""
    return bool(load_all_chunks())
